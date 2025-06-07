# -*- coding: utf-8 -*-
"""
Archivo: Exporter.py
Descripción: Módulo para gestionar la exportación de datos de libros a diferentes formatos (TXT, DOCX, PDF).
Autor: AutoDoc AI
Fecha: 07/06/2025
Versión: 0.0.1
Licencia: MIT License
"""

import os
import re
import html
from typing import Dict, List, Any, Optional
import traceback

# --- Detección de Librería DOCX ---
# Verifica si la librería 'python-docx' está instalada y disponible.
PYTHON_DOCX_AVAILABLE = False
try:
    # Intenta importar las clases necesarias de python-docx
    from docx import Document
    from docx.shared import Inches, Pt as DocxPt
    from docx.enum.text import WD_ALIGN_PARAGRAPH as DocxWD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
    print("INFO (Exporter.py): Librería 'python-docx' cargada exitosamente.")
except Exception as e_docx_import:
    # Si falla la importación, imprime un error y desactiva la funcionalidad DOCX.
    print("--------------------------------------------------------------------")
    print("ERROR (Exporter.py): Falló la importación de 'python-docx'.")
    print(f"Detalle del error: {e_docx_import}")
    traceback.print_exc()
    print("La exportación a DOCX NO estará disponible.")
    print("Asegúrese de que 'python-docx' está instalado: pip install python-docx")
    print("--------------------------------------------------------------------")
    # Define clases stub para evitar errores de NameError si la importación falla.
    class Document:
        """Clase stub para `docx.Document` cuando la librería no está disponible."""
        pass
    class Inches:
        """Clase stub para `docx.shared.Inches` cuando la librería no está disponible."""
        pass
    class DocxPt:
        """Clase stub para `docx.shared.Pt` cuando la librería no está disponible."""
        pass
    class DocxWD_ALIGN_PARAGRAPH:
        """Clase stub para `docx.enum.text.WD_ALIGN_PARAGRAPH` cuando la librería no está disponible."""
        pass

# --- Detección de Librería PDF (ReportLab) ---
# Verifica si la librería 'reportlab' está instalada y disponible.
REPORTLAB_AVAILABLE = False
try:
    # Intenta importar las clases y constantes necesarias de reportlab.
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as ReportLabImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.lib.units import inch, cm
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.colors import black, lightgrey, HexColor
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    REPORTLAB_AVAILABLE = True
    print("INFO (Exporter.py): Librería 'reportlab' cargada exitosamente.")
except Exception as e_reportlab_import:
    # Si falla la importación, imprime un error y desactiva la funcionalidad PDF.
    print("--------------------------------------------------------------------")
    print("ERROR (Exporter.py): Falló la importación de 'reportlab'.")
    print(f"Detalle del error: {e_reportlab_import}")
    traceback.print_exc()
    print("La exportación a PDF NO estará disponible.")
    print("Asegúrese de que 'reportlab' está instalado: pip install reportlab")
    print("--------------------------------------------------------------------")
    # Define clases stub para evitar errores de NameError si la importación falla.
    class SimpleDocTemplate:
        """Clase stub para `reportlab.platypus.SimpleDocTemplate`."""
        pass
    class Paragraph:
        """Clase stub para `reportlab.platypus.Paragraph`."""
        pass
    class Spacer:
        """Clase stub para `reportlab.platypus.Spacer`."""
        pass
    class PageBreak:
        """Clase stub para `reportlab.platypus.PageBreak`."""
        pass
    class ReportLabImage:
        """Clase stub para `reportlab.platypus.Image`."""
        pass
    class getSampleStyleSheet:
        """Clase stub para `reportlab.lib.styles.getSampleStyleSheet`."""
        pass
    class ParagraphStyle:
        """Clase stub para `reportlab.lib.styles.ParagraphStyle`."""
        pass
    class TA_CENTER:
        """Clase stub para `reportlab.lib.enums.TA_CENTER`."""
        pass
    class TA_JUSTIFY:
        """Clase stub para `reportlab.lib.enums.TA_JUSTIFY`."""
        pass
    class TA_LEFT:
        """Clase stub para `reportlab.lib.enums.TA_LEFT`."""
        pass
    class inch:
        """Clase stub para `reportlab.lib.units.inch`."""
        pass
    class cm:
        """Clase stub para `reportlab.lib.units.cm`."""
        pass
    class A4:
        """Clase stub para `reportlab.lib.pagesizes.A4`."""
        pass
    class black:
        """Clase stub para `reportlab.lib.colors.black`."""
        pass
    class lightgrey:
        """Clase stub para `reportlab.lib.colors.lightgrey`."""
        pass
    class HexColor:
        """Clase stub para `reportlab.lib.colors.HexColor`."""
        pass
    class canvas:
        """Clase stub para `reportlab.pdfgen.canvas`."""
        pass
    class ImageReader:
        """Clase stub para `reportlab.lib.utils.ImageReader`."""
        pass


class BaseExporter:
    """
    Clase base abstracta para los exportadores de formatos de libro.

    Define la interfaz común que deben implementar todas las clases exportadoras
    específicas de formato. También incluye métodos de utilidad para el manejo
    de contenido HTML.
    """
    def export(self, book_data: Dict[str, Any], chapters_data: List[Dict[str, Any]], output_path: str) -> bool:
        """
        Método principal de exportación (abstracto).

        Este método debe ser implementado por las subclases para realizar la
        exportación real a un formato específico.

        Args:
            book_data (Dict[str, Any]): Diccionario con los metadatos del libro
                                         (título, autor, prólogo, ruta de imagen de portada, etc.).
            chapters_data (List[Dict[str, Any]]): Lista de diccionarios, donde cada uno
                                                  representa un capítulo (número, título, contenido, etc.).
            output_path (str): La ruta completa del archivo de salida donde se guardará
                               el libro exportado.

        Returns:
            bool: True si la exportación fue exitosa, False en caso contrario.

        Raises:
            NotImplementedError: Siempre se lanza en la clase base para indicar que
                                 el método debe ser sobrescrito.
        """
        raise NotImplementedError("El método 'export' debe ser implementado por las subclases.")

    def _clean_html_for_plaintext(self, html_content: Optional[str]) -> str:
        """
        Limpia una cadena de contenido HTML para obtener una representación de texto plano.

        Elimina la mayoría de las etiquetas HTML, reemplaza `<br>` con saltos de línea
        y decodifica entidades HTML.

        Args:
            html_content (Optional[str]): La cadena de contenido HTML a limpiar. Puede ser None.

        Returns:
            str: La cadena resultante con solo texto plano. Retorna una cadena vacía
                 si la entrada es None o vacía.
        """
        if not html_content:
            return ""
        # Reemplaza las etiquetas <br> con saltos de línea.
        text = re.sub(r'<br\s*/?>', '\n', html_content, flags=re.IGNORECASE)
        # Elimina todas las demás etiquetas HTML.
        text = re.sub(r'<[^>]+>', '', text)
        # Decodifica entidades HTML (como &nbsp;, &amp;, etc.).
        text = html.unescape(text)
        # Elimina espacios en blanco al inicio y final.
        return text.strip()

    def _transform_html_for_reportlab(self, html_content: Optional[str]) -> str:
        """
        Transforma el HTML de entrada al subconjunto limitado que ReportLab Paragraph entiende.

        ReportLab Paragraphs soporta un subconjunto de etiquetas HTML como <b>, <i>, <u>,
        <br/>, <font color="..." size="..." name="...">. Este método adapta el HTML
        original para ser compatible, prestando especial atención a la etiqueta <font>
        y las entidades HTML.

        Args:
            html_content (Optional[str]): La cadena de contenido HTML a transformar. Puede ser None.

        Returns:
            str: La cadena HTML transformada, compatible con ReportLab Paragraph. Retorna
                 una cadena vacía si la entrada es None o vacía.
        """
        if not html_content:
            return ""

        text = str(html_content)
        # Asegura que <br> sea <br/> para ReportLab.
        text = re.sub(r'<br\s*/?>', '<br/>', text, flags=re.IGNORECASE)
        # Reemplaza la etiqueta <font> usando el método estático auxiliar.
        text = re.sub(r'<font([^>]*)>', PdfExporter._replace_font_tag_for_reportlab_static, text, flags=re.IGNORECASE)
        # Maneja entidades HTML comunes que ReportLab podría no decodificar automáticamente.
        text = text.replace('&', '&') # Esto es un truco común para asegurar que &amp; no se doble decodifique.
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        text = text.replace('&amp;', '&')
        text = text.replace('&nbsp;', ' ') # Reemplaza &nbsp; con un espacio no rompible Unicode.
        return text

class TxtExporter(BaseExporter):
    """
    Exportador para generar archivos de texto plano (.txt).

    Implementa la lógica para tomar los datos de un libro y sus capítulos y
    escribirlos en un archivo de texto plano, eliminando todo el formato HTML.
    """
    def export(self, book_data: Dict[str, Any], chapters_data: List[Dict[str, Any]], output_path: str) -> bool:
        """
        Implementa la exportación del libro a un archivo de texto plano.

        Genera un archivo .txt con el título, autor, un índice simple, el prólogo
        y el contenido de cada capítulo, todo en texto plano.

        Args:
            book_data (Dict[str, Any]): Diccionario con los metadatos del libro.
            chapters_data (List[Dict[str, Any]]): Lista de diccionarios de capítulos.
            output_path (str): La ruta completa del archivo .txt de salida.

        Returns:
            bool: True si la exportación fue exitosa, False en caso contrario.

        Raises:
            IOError: Si ocurre un error al escribir en el archivo.
            Exception: Para otros errores inesperados durante la exportación.
        """
        try:
            # Abre el archivo de salida en modo escritura con codificación UTF-8.
            with open(output_path, 'w', encoding='utf-8') as f:
                # Escribe el título del libro.
                title = book_data.get('title', 'Libro Sin Título').upper()
                f.write(f"{title}\n")
                f.write("=" * len(title) + "\n\n")
                # Escribe el autor.
                author = book_data.get('author', 'Autor Desconocido')
                f.write(f"Por: {author}\n\n")
                f.write("-" * 30 + "\n\n")

                # Escribe el índice de capítulos.
                f.write("ÍNDICE\n")
                f.write("------\n")
                if not chapters_data:
                    f.write("(No hay capítulos)\n")
                else:
                    # Ordena los capítulos por número antes de listarlos.
                    sorted_chapters = sorted(chapters_data, key=lambda chp: chp.get('chapter_number', float('inf')))
                    for chapter in sorted_chapters:
                        ch_num = chapter.get('chapter_number', '#')
                        # Limpia el título del capítulo de HTML.
                        ch_title = self._clean_html_for_plaintext(chapter.get('title', 'Capítulo Sin Título'))
                        f.write(f"Capítulo {ch_num}: {ch_title}\n")
                f.write("\n" + "-" * 30 + "\n\n")

                # Escribe el prólogo si existe.
                prologue_content_html = book_data.get('prologue', '')
                if prologue_content_html:
                    f.write("PRÓLOGO\n")
                    f.write("---------\n")
                    # Limpia el contenido del prólogo de HTML.
                    prologue_plain = self._clean_html_for_plaintext(prologue_content_html)
                    f.write(prologue_plain + "\n\n")
                    f.write("-" * 30 + "\n\n")

                # Escribe el contenido de cada capítulo.
                if not chapters_data:
                    f.write("No hay contenido de capítulos para exportar.\n")
                else:
                    # Ordena los capítulos por número antes de escribir su contenido.
                    sorted_chapters = sorted(chapters_data, key=lambda chp: chp.get('chapter_number', float('inf')))
                    for chapter in sorted_chapters:
                        ch_num = chapter.get('chapter_number', '#')
                        # Limpia el título del capítulo para el encabezado.
                        ch_title_plain = self._clean_html_for_plaintext(chapter.get('title', 'Capítulo Sin Título'))
                        header = f"CAPÍTULO {ch_num}: {ch_title_plain}"
                        f.write(f"{header.upper()}\n")
                        f.write("-" * len(header) + "\n")
                        content_html = chapter.get('content', '')
                        # Limpia el contenido del capítulo de HTML.
                        content_plain = self._clean_html_for_plaintext(content_html)
                        f.write(content_plain + "\n\n")

            # Retorna True si todo se escribió sin errores.
            return True
        except IOError as e:
            # Captura y reporta errores de entrada/salida.
            print(f"Error de E/S al escribir el archivo TXT '{output_path}': {e}")
            return False
        except Exception as e_gen:
            # Captura y reporta otros errores generales.
            print(f"Error inesperado durante la exportación TXT a '{output_path}': {e_gen}")
            return False

class DocxExporter(BaseExporter):
    """
    Exportador para generar archivos Microsoft Word (.docx).

    Utiliza la librería `python-docx` para crear un documento Word con el contenido
    del libro, incluyendo metadatos, índice, prólogo y capítulos. Intenta preservar
    algo de formato básico y añadir la imagen de portada.
    """
    def export(self, book_data: Dict[str, Any], chapters_data: List[Dict[str, Any]], output_path: str) -> bool:
        """
        Implementa la exportación del libro a un archivo DOCX.

        Crea un documento .docx con la imagen de portada (si existe), título, autor,
        índice, prólogo y contenido de los capítulos. El contenido HTML se convierte
        a texto plano antes de insertarlo.

        Args:
            book_data (Dict[str, Any]): Diccionario con los metadatos del libro.
            chapters_data (List[Dict[str, Any]]): Lista de diccionarios de capítulos.
            output_path (str): La ruta completa del archivo .docx de salida.

        Returns:
            bool: True si la exportación fue exitosa, False en caso contrario. Retorna
                  False inmediatamente si la librería `python-docx` no está disponible.

        Raises:
            IOError: Si ocurre un error al escribir en el archivo.
            Exception: Para otros errores inesperados durante la exportación.
        """
        # Verifica si la librería python-docx está disponible.
        if not PYTHON_DOCX_AVAILABLE:
            print("Error: La librería 'python-docx' no está disponible para exportar a DOCX.")
            return False

        try:
            # Crea un nuevo documento Word.
            document = Document()

            # Intenta añadir la imagen de portada.
            cover_image_path = book_data.get('cover_image_path')
            if cover_image_path and os.path.exists(cover_image_path):
                try:
                    # Añade la imagen y la centra.
                    document.add_picture(cover_image_path, width=Inches(5.0))
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = DocxWD_ALIGN_PARAGRAPH.CENTER
                    document.add_paragraph() # Añade un espacio después de la imagen.
                except Exception as e_img:
                    print(f"Advertencia: No se pudo añadir la imagen de portada '{cover_image_path}' al DOCX: {e_img}")

            # Añade el título del libro.
            title_text = book_data.get('title', 'Libro Sin Título')
            title_paragraph = document.add_paragraph()
            title_run = title_paragraph.add_run(title_text)
            # Aplica formato básico al título.
            title_run.font.name = 'Calibri'
            title_run.font.size = DocxPt(18)
            title_run.bold = True
            title_paragraph.alignment = DocxWD_ALIGN_PARAGRAPH.CENTER

            # Añade el autor.
            author_text = book_data.get('author', 'Autor Desconocido')
            author_paragraph = document.add_paragraph()
            author_run = author_paragraph.add_run(f"Por: {author_text}")
            # Aplica formato básico al autor.
            author_run.font.name = 'Calibri'
            author_run.font.size = DocxPt(10)
            author_paragraph.alignment = DocxWD_ALIGN_PARAGRAPH.CENTER

            # Añade un salto de página antes del índice.
            document.add_page_break()

            # Añade el título del índice.
            index_title_paragraph = document.add_paragraph()
            index_title_run = index_title_paragraph.add_run("ÍNDICE")
            index_title_run.font.size = DocxPt(14)
            index_title_run.bold = True

            # Añade los elementos del índice.
            if not chapters_data:
                document.add_paragraph("(No hay capítulos)")
            else:
                # Ordena los capítulos por número.
                sorted_chapters = sorted(chapters_data, key=lambda chp: chp.get('chapter_number', float('inf')))
                for chapter in sorted_chapters:
                    ch_num = chapter.get('chapter_number', '#')
                    # Limpia el título del capítulo de HTML para el índice.
                    ch_title = self._clean_html_for_plaintext(chapter.get('title', 'Capítulo Sin Título'))
                    # Añade cada capítulo como un elemento de lista con viñetas.
                    document.add_paragraph(f"Capítulo {ch_num}: {ch_title}", style='ListBullet')

            # Añade el prólogo si existe.
            prologue_content_html = book_data.get('prologue', '')
            if prologue_content_html:
                # Añade un salto de página antes del prólogo.
                document.add_page_break()
                # Añade el título del prólogo.
                prologue_title_paragraph = document.add_paragraph()
                prologue_title_run = prologue_title_paragraph.add_run("PRÓLOGO")
                prologue_title_run.font.size = DocxPt(14)
                prologue_title_run.bold = True
                # Limpia el contenido del prólogo de HTML y lo añade.
                prologue_plain = self._clean_html_for_plaintext(prologue_content_html)
                document.add_paragraph(prologue_plain)

            # Añade el contenido de cada capítulo.
            if chapters_data:
                # Ordena los capítulos por número.
                sorted_chapters = sorted(chapters_data, key=lambda chp: chp.get('chapter_number', float('inf')))
                for chapter in sorted_chapters:
                    # Añade un salto de página antes de cada capítulo.
                    document.add_page_break()
                    ch_num = chapter.get('chapter_number', '#')
                    # Limpia el título del capítulo para el encabezado.
                    ch_title_plain = self._clean_html_for_plaintext(chapter.get('title', 'Capítulo Sin Título'))
                    # Añade el encabezado del capítulo.
                    chapter_header_paragraph = document.add_paragraph()
                    chapter_header_run = chapter_header_paragraph.add_run(f"CAPÍTULO {ch_num}: {ch_title_plain.upper()}")
                    chapter_header_run.font.size = DocxPt(14)
                    chapter_header_run.bold = True
                    # Limpia el contenido del capítulo de HTML y lo añade.
                    content_html = chapter.get('content', '')
                    content_plain = self._clean_html_for_plaintext(content_html)
                    document.add_paragraph(content_plain)
            # Si no hay capítulos ni prólogo, añade un mensaje.
            elif not prologue_content_html:
                 document.add_page_break()
                 document.add_paragraph("No hay contenido de capítulos para exportar.")

            # Guarda el documento en la ruta especificada.
            document.save(output_path)
            return True
        except IOError as e:
            # Captura y reporta errores de entrada/salida.
            print(f"Error de E/S al escribir el archivo DOCX '{output_path}': {e}")
            return False
        except Exception as e_gen:
            # Captura y reporta otros errores generales.
            print(f"Error inesperado durante la exportación DOCX a '{output_path}': {e_gen}")
            return False

class PdfExporter(BaseExporter):
    """
    Exportador para generar archivos PDF usando la librería ReportLab.

    Permite crear un documento PDF con una página de portada (opcionalmente
    con imagen de fondo), índice, prólogo y capítulos. Intenta preservar
    un subconjunto limitado de formato HTML.
    """
    # Variables de clase para pasar datos a la función de plantilla de página.
    _first_page_rendered: bool = False
    _book_data_for_first_page: Optional[Dict[str, Any]] = None

    @staticmethod
    def _replace_font_tag_for_reportlab_static(match_obj: re.Match) -> str:
        """
        Método estático auxiliar para reemplazar etiquetas <font> por un formato compatible con ReportLab.

        ReportLab Paragraphs entiende <font name="..." size="..." color="...">.
        Este método toma los atributos `data-point-size` y `color` del HTML original
        y los transforma a los atributos `size` y `color` esperados por ReportLab.
        Establece el nombre de la fuente a "Helvetica" por defecto.

        Args:
            match_obj (re.Match): El objeto de coincidencia del regex para la etiqueta <font>.

        Returns:
            str: La cadena de la etiqueta <font> transformada para ReportLab.
        """
        attrs_str = match_obj.group(1) # Captura los atributos dentro de la etiqueta font.
        # Busca el tamaño de fuente en el atributo data-point-size.
        size_match = re.search(r'data-point-size="(\d+)"', attrs_str, re.IGNORECASE)
        # Busca el color en el atributo color.
        color_match = re.search(r'color="(#[0-9a-fA-F]{6})"', attrs_str, re.IGNORECASE)

        new_attrs_list = ['name="Helvetica"'] # Define la fuente por defecto.
        if size_match:
            # Añade el tamaño si se encontró.
            new_attrs_list.append(f'size="{size_match.group(1)}"')
        if color_match:
            # Añade el color si se encontró.
            new_attrs_list.append(f'color="{color_match.group(1)}"')

        # Une los atributos para formar la nueva etiqueta <font>.
        new_font_attrs_str = ' '.join(new_attrs_list)
        return f"<font {new_font_attrs_str}>"

    def _draw_cover_page_template(self, canvas_obj: canvas.Canvas, doc: SimpleDocTemplate):
        """
        Dibuja la imagen de portada como fondo y el texto superpuesto en la primera página.

        Esta función se pasa a `doc.build` como `onFirstPage` y `onLaterPages`.
        Solo dibuja la portada en la primera página y solo una vez.

        Args:
            canvas_obj (canvas.Canvas): El objeto canvas de ReportLab para dibujar.
            doc (SimpleDocTemplate): El objeto documento de ReportLab.
        """
        # Solo ejecuta en la primera página y si aún no se ha dibujado.
        if doc.page == 1 and not PdfExporter._first_page_rendered:
            canvas_obj.saveState() # Guarda el estado actual del canvas.
            page_width, page_height = A4 # Obtiene las dimensiones de la página A4.

            # Intenta dibujar la imagen de portada como fondo.
            if PdfExporter._book_data_for_first_page:
                cover_image_path = PdfExporter._book_data_for_first_page.get('cover_image_path')
                if cover_image_path and os.path.exists(cover_image_path):
                    try:
                        img_reader = ImageReader(cover_image_path)
                        img_width, img_height = img_reader.getSize()
                        img_aspect_ratio = img_height / float(img_width)

                        # Calcula las dimensiones para que la imagen cubra la página manteniendo el aspecto.
                        target_w = page_width
                        target_h = target_w * img_aspect_ratio
                        if target_h < page_height:
                            target_h = page_height
                            target_w = target_h / img_aspect_ratio

                        # Calcula la posición para centrar la imagen.
                        x_offset = (page_width - target_w) / 2.0
                        y_offset = (page_height - target_h) / 2.0

                        # Dibuja la imagen.
                        canvas_obj.drawImage(img_reader, x_offset, y_offset, width=target_w, height=target_h, mask='auto')
                    except Exception as e_img:
                        print(f"Advertencia (PDF): No se pudo dibujar la imagen de portada '{cover_image_path}': {e_img}")

            # Dibuja el título y autor superpuestos.
            if PdfExporter._book_data_for_first_page:
                # Limpia el título y autor de HTML.
                title_text = self._clean_html_for_plaintext(PdfExporter._book_data_for_first_page.get('title', 'Libro Sin Título'))
                # Configura el color y la fuente para el texto.
                canvas_obj.setFillColor(black)
                canvas_obj.setFont("Helvetica-Bold", 18)
                # Dibuja el título centrado.
                canvas_obj.drawCentredString(page_width / 2.0, page_height - 3*inch, title_text)

                author_text = self._clean_html_for_plaintext(PdfExporter._book_data_for_first_page.get('author', 'Autor Desconocido'))
                canvas_obj.setFont("Helvetica", 10)
                author_full_text = f"Por: {author_text}"
                # Dibuja el autor centrado.
                canvas_obj.drawCentredString(page_width / 2.0, page_height - 3.5*inch, author_full_text)

            canvas_obj.restoreState() # Restaura el estado del canvas.
            PdfExporter._first_page_rendered = True # Marca la primera página como dibujada.

    def export(self, book_data: Dict[str, Any], chapters_data: List[Dict[str, Any]], output_path: str) -> bool:
        """
        Implementa la exportación del libro a un archivo PDF usando ReportLab.

        Crea un documento PDF con una página de portada (si hay imagen), índice,
        prólogo y contenido de los capítulos. El contenido HTML se transforma
        para ser compatible con ReportLab Paragraphs.

        Args:
            book_data (Dict[str, Any]): Diccionario con los metadatos del libro.
            chapters_data (List[Dict[str, Any]]): Lista de diccionarios de capítulos.
            output_path (str): La ruta completa del archivo .pdf de salida.

        Returns:
            bool: True si la exportación fue exitosa, False en caso contrario. Retorna
                  False inmediatamente si la librería `reportlab` no está disponible.

        Raises:
            IOError: Si ocurre un error al escribir en el archivo.
            Exception: Para otros errores inesperados durante la exportación.
        """
        # Verifica si la librería reportlab está disponible.
        if not REPORTLAB_AVAILABLE:
            print("Error: La librería 'reportlab' no está disponible para exportar a PDF.")
            return False

        # Almacena los datos del libro para que la plantilla de página pueda acceder a ellos.
        PdfExporter._book_data_for_first_page = book_data
        PdfExporter._first_page_rendered = False # Reinicia el indicador de página dibujada.

        try:
            # Configura el documento PDF.
            doc = SimpleDocTemplate(output_path, pagesize=A4,
                                    rightMargin=1.5*cm, leftMargin=1.5*cm,
                                    topMargin=2*cm, bottomMargin=2*cm)

            # Obtiene estilos de párrafo de ejemplo y define estilos personalizados.
            styles = getSampleStyleSheet()
            style_section_title = ParagraphStyle('SectionTitle', parent=styles['h2'],
                                           fontSize=14, leading=18, spaceBefore=12, spaceAfter=6,
                                           fontName='Helvetica-Bold', alignment=TA_LEFT)
            style_body = ParagraphStyle('BodyText', parent=styles['Normal'],
                                      alignment=TA_JUSTIFY, spaceBefore=0, spaceAfter=6,
                                      fontName='Helvetica', fontSize=10, leading=14,
                                      firstLineIndent=0.5*cm)
            style_index_item = ParagraphStyle('IndexItem', parent=styles['Normal'],
                                            leftIndent=0.5*cm, fontName='Helvetica', fontSize=10, leading=12)

            # Lista de elementos (Story) que se añadirán al documento.
            story: List[Any] = []

            # Añade un salto de página para que la primera página sea la de portada.
            story.append(PageBreak())

            # Añade el título del índice.
            story.append(Paragraph("ÍNDICE", style_section_title))
            story.append(Spacer(1, 0.2*cm)) # Añade un espacio.

            # Añade los elementos del índice.
            if not chapters_data:
                story.append(Paragraph("(No hay capítulos)", style_body))
            else:
                # Ordena los capítulos por número.
                sorted_chapters = sorted(chapters_data, key=lambda chp: chp.get('chapter_number', float('inf')))
                for chapter in sorted_chapters:
                    ch_num = chapter.get('chapter_number', '#')
                    ch_title_html = chapter.get('title', 'Capítulo Sin Título')
                    # Transforma el HTML del título para ReportLab.
                    ch_title_transformed = self._transform_html_for_reportlab(ch_title_html)
                    # Añade el elemento del índice como un párrafo con estilo.
                    story.append(Paragraph(f"Capítulo {ch_num}: {ch_title_transformed}", style_index_item))

            story.append(Spacer(1, 0.5*cm)) # Añade un espacio.

            # Añade el prólogo si existe.
            prologue_html = book_data.get('prologue', '')
            if prologue_html:
                story.append(PageBreak()) # Salto de página antes del prólogo.
                story.append(Paragraph("PRÓLOGO", style_section_title)) # Título del prólogo.
                # Transforma el HTML del prólogo y lo añade como párrafo.
                prologue_transformed = self._transform_html_for_reportlab(prologue_html)
                story.append(Paragraph(prologue_transformed, style_body))

            # Añade el contenido de cada capítulo.
            if chapters_data:
                # Ordena los capítulos por número.
                sorted_chapters = sorted(chapters_data, key=lambda chp: chp.get('chapter_number', float('inf')))
                for chapter in sorted_chapters:
                    story.append(PageBreak()) # Salto de página antes de cada capítulo.
                    ch_num = chapter.get('chapter_number', '#')
                    ch_title_html = chapter.get('title', 'Capítulo Sin Título')
                    # Transforma el HTML del título del capítulo.
                    ch_title_transformed = self._transform_html_for_reportlab(ch_title_html)
                    # Añade el encabezado del capítulo.
                    story.append(Paragraph(f"CAPÍTULO {ch_num}: {ch_title_transformed.upper()}", style_section_title))
                    content_html = chapter.get('content', '')
                    # Transforma el HTML del contenido del capítulo y lo añade como párrafo.
                    content_transformed = self._transform_html_for_reportlab(content_html)
                    story.append(Paragraph(content_transformed, style_body))
            # Si no hay capítulos ni prólogo, añade un mensaje.
            elif not prologue_html:
                story.append(PageBreak())
                story.append(Paragraph("No hay contenido de capítulos para exportar.", style_body))

            # Construye el documento PDF, usando la función de plantilla para la primera página.
            doc.build(story, onFirstPage=self._draw_cover_page_template, onLaterPages=self._draw_cover_page_template)

            return True
        except IOError as e:
            # Captura y reporta errores de entrada/salida.
            print(f"Error de E/S al escribir el archivo PDF '{output_path}': {e}")
            return False
        except Exception as e_gen:
            # Captura y reporta otros errores generales.
            print(f"Error inesperado durante la exportación PDF a '{output_path}': {e_gen}")
            traceback.print_exc() # Imprime el traceback completo para depuración.
            return False
        finally:
            # Limpia las variables de clase después de la exportación.
            PdfExporter._book_data_for_first_page = None
            PdfExporter._first_page_rendered = False